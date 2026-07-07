#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Costruisce TecnologIA_Visual_Brief_ESPLORA.docx consolidando i 56 brief .md
generati da _generate_visual_briefs.py.

Output: 04_CONTENUTI/visual_esplora/TecnologIA_Visual_Brief_ESPLORA.docx
"""

from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ────────────────────────────────────────────────────────────────────
# CONFIG
# ────────────────────────────────────────────────────────────────────

OUT_FILE = "TecnologIA_Visual_Brief_ESPLORA.docx"

AREA_COLORS = {
    "MAT": "6D4C41",
    "DIS": "1565C0",
    "DIG": "006064",
    "ALI": "558B2F",
    "AMB": "00695C",
    "ENE": "E65100",
    "COM": "283593",
    "SIS": "4527A0",
    "INF": "0277BD",
}

AREA_LABELS = {
    "MAT": "🪨 Materiali e Rifiuti",
    "DIS": "📐 Disegno Tecnico",
    "DIG": "💻 Digitale / Coding / AI",
    "ALI": "🌾 Alimentazione",
    "AMB": "🏗️ Abitazione · Città · Territorio",
    "ENE": "⚡ Energia e Macchine",
    "COM": "📡 Comunicazioni e Trasporti",
    "SIS": "⚙️ Sistemi · Economia · Lavoro",
    "INF": "🖥️ Informatica",
}


# ────────────────────────────────────────────────────────────────────
# UTILITY DOCX
# ────────────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_horizontal_line(paragraph):
    """Aggiunge una linea orizzontale come bordo inferiore del paragrafo."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_styles(doc: Document):
    """Imposta gli stili base del documento."""
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    for level, size in [(1, 22), (2, 16), (3, 13), (4, 11)]:
        st = styles[f"Heading {level}"]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


# ────────────────────────────────────────────────────────────────────
# PARSER MARKDOWN BRIEF
# ────────────────────────────────────────────────────────────────────

def parse_brief_md(text: str) -> dict:
    """Estrae i campi salienti da un brief .md."""
    out: dict = {}

    # Titoli H1 (MC id) e H2 (titolo MC)
    m = re.search(r"^#\s+Visual Brief ESPLORA\s*[—–-]\s*(MC-\S+)$", text, re.M)
    if m:
        out["mc_id"] = m.group(1)
    m2 = re.search(r"^##\s+(.+)$", text, re.M)
    if m2:
        out["mc_titolo"] = m2.group(1).strip()

    # Tabella meta in testa
    meta = {}
    table_match = re.search(r"\| Campo \| Valore \|\n\|[^\n]+\|\n((?:\|[^\n]+\|\n?)+)", text)
    if table_match:
        for row in table_match.group(1).strip().split("\n"):
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            if len(cells) == 2:
                key = re.sub(r"\*\*", "", cells[0]).strip()
                val = re.sub(r"`", "", cells[1]).strip()
                meta[key] = val
    out["meta"] = meta

    # Immagini
    images = []
    # Split per "## Immagine NN — ..."
    img_blocks = re.split(r"\n## Immagine\s+(\d+)\s+[—–-]\s+", text)
    # img_blocks[0] è il preambolo, poi alternati: numero, contenuto
    for i in range(1, len(img_blocks), 2):
        n = img_blocks[i]
        body = img_blocks[i + 1] if i + 1 < len(img_blocks) else ""
        # Titolo immagine fino al newline
        first_line, _, rest = body.partition("\n")
        img: dict = {
            "n": int(n),
            "titolo": first_line.strip(),
        }
        # Tabella meta immagine
        meta_img = {}
        t2 = re.search(r"\| Campo \| Valore \|\n\|[^\n]+\|\n((?:\|[^\n]+\|\n?)+)", rest)
        if t2:
            for row in t2.group(1).strip().split("\n"):
                cells = [c.strip() for c in row.strip().strip("|").split("|")]
                if len(cells) == 2:
                    key = re.sub(r"\*\*", "", cells[0]).strip()
                    val = cells[1].strip()
                    meta_img[key] = val
        img["meta"] = meta_img

        # Sotto-sezioni interne (eventuale blockquote)
        sub = re.search(r"^>\s*_Sotto-sezioni interne:\s*(.+?)_\s*$", rest, re.M)
        img["sotto_sezioni"] = sub.group(1).strip() if sub else None

        # Testo di riferimento
        ref = re.search(r"\*\*Testo di riferimento:\*\*\s*\n>\s*(.+?)(?=\n\n|\n\*\*Prompt)", rest, re.S)
        img["testo_riferimento"] = ref.group(1).strip() if ref else ""

        # Prompt (dentro ```)
        pr = re.search(r"\*\*Prompt:\*\*\s*\n```\s*\n(.+?)\n```", rest, re.S)
        img["prompt"] = pr.group(1).strip() if pr else ""

        images.append(img)

    out["images"] = images
    return out


# ────────────────────────────────────────────────────────────────────
# RENDERING DOCX
# ────────────────────────────────────────────────────────────────────

def add_brief_section(doc: Document, brief: dict, area_color: str):
    """Aggiunge una sezione documento per un singolo MC brief."""
    # Page break prima di ogni MC tranne il primo (gestito dal chiamante)
    # Heading 1: MC id + titolo
    h1 = doc.add_paragraph(style="Heading 1")
    run = h1.add_run(f"{brief['mc_id']} — {brief.get('mc_titolo', '')}")
    run.font.color.rgb = RGBColor.from_string(area_color)

    # Tabella metadata MC
    meta = brief.get("meta", {})
    if meta:
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Light List Accent 1"
        for key in ["Area", "Anno", "Livello DigComp", "Colore area", "Generato", "Immagini totali"]:
            if key in meta:
                row = tbl.add_row().cells
                row[0].text = key
                row[1].text = meta[key]
                for c in row:
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)
                set_cell_shading(row[0], "F0F0F0")

    doc.add_paragraph()

    # Per ogni immagine
    for img in brief.get("images", []):
        # Heading 2: numero immagine + titolo sezione
        h2 = doc.add_paragraph(style="Heading 2")
        h2_run = h2.add_run(f"Immagine {img['n']:02d} — {img['titolo']}")
        h2_run.font.color.rgb = RGBColor.from_string(area_color)

        # Tabella meta immagine (Tipo, Formato, Legenda, Posizione, Etichette curate)
        meta_img = img.get("meta", {})
        if meta_img:
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Light Grid Accent 1"
            tbl.columns[0].width = Cm(4)
            tbl.columns[1].width = Cm(12)
            for key in ["Tipo", "Formato", "Legenda", "Posizione", "Etichette curate"]:
                if key in meta_img:
                    row = tbl.add_row().cells
                    row[0].text = key
                    row[1].text = meta_img[key]
                    for c in row:
                        for p in c.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(9)
                    set_cell_shading(row[0], "F8F8F8")

        # Sotto-sezioni interne (se presenti)
        if img.get("sotto_sezioni"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            run = p.add_run(f"Sotto-sezioni interne: {img['sotto_sezioni']}")
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Testo di riferimento
        if img.get("testo_riferimento"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.add_run("Testo di riferimento: ").bold = True
            run = p.add_run(img["testo_riferimento"])
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            p.paragraph_format.left_indent = Cm(0.5)

        # Prompt — in box monospace
        if img.get("prompt"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run("Prompt per image-gen:")
            run.bold = True

            # Box prompt
            prompt_tbl = doc.add_table(rows=1, cols=1)
            prompt_tbl.style = "Table Grid"
            cell = prompt_tbl.rows[0].cells[0]
            set_cell_shading(cell, "FAFAFA")
            cell.text = ""
            cp = cell.paragraphs[0]
            cp_run = cp.add_run(img["prompt"])
            cp_run.font.name = "Consolas"
            cp_run.font.size = Pt(8.5)
            cp_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            # Bordi sottili
            cell_tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:color"), "CCCCCC")
                tc_borders.append(b)
            cell_tc_pr.append(tc_borders)

        # Separatore sotto ogni immagine
        sep = doc.add_paragraph()
        add_horizontal_line(sep)
        sep.paragraph_format.space_after = Pt(6)


def build_cover(doc: Document, n_brief: int, n_img: int):
    """Pagina copertina."""
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(80)
    run = t.add_run("TecnologIA")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = t2.add_run("Visual Brief ESPLORA — Volume consolidato")
    run2.font.size = Pt(20)
    run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.paragraph_format.space_before = Pt(40)
    run3 = t3.add_run(f"56 micro-competenze · {n_img} immagini · prompt v2 descrittivi di scena")
    run3.font.size = Pt(12)
    run3.font.italic = True
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Disclaimer prompt v2
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(60)
    note.paragraph_format.left_indent = Cm(2)
    note.paragraph_format.right_indent = Cm(2)
    nrun = note.add_run(
        "Questo volume contiene i brief operativi per la generazione delle immagini ESPLORA. "
        "I prompt sono in inglese (ottimizzati per modelli image-gen come GPT Image 2, FLUX, Midjourney v6+) "
        "e specificano le etichette da rendere in italiano dentro le immagini.\n\n"
        "I prompt sono descrittivi di scena: includono personaggio, ambientazione, oggetti specifici dell'area "
        "didattica, palette di colori e regole tipografiche. Il titolo italiano della sezione non va mai "
        "renderizzato nell'immagine come testo.\n\n"
        "Riferimento operativo: 07_GUIDE/operative/PROMPT_PATTERNS_visual_ESPLORA.md\n"
        "Generatore: 04_CONTENUTI/visual_esplora/_generate_visual_briefs.py"
    )
    nrun.font.size = Pt(10)
    nrun.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.space_before = Pt(60)
    today = date.today().isoformat()
    frun = foot.add_run(f"Generato: {today}  ·  Autore: Antonio Scaramuzzino")
    frun.font.size = Pt(9)
    frun.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Page break
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_index(doc: Document, briefs: list[dict]):
    """Indice dei brief raggruppati per anno+area."""
    h = doc.add_paragraph(style="Heading 1")
    h.add_run("Indice dei brief").font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    by_anno_area: dict = {}
    for b in briefs:
        m = re.match(r"MC-([A-Z]+)-(\d)-\d+", b["mc_id"])
        if not m:
            continue
        area = m.group(1)
        anno = int(m.group(2))
        by_anno_area.setdefault((anno, area), []).append(b)

    for (anno, area) in sorted(by_anno_area):
        h2 = doc.add_paragraph(style="Heading 3")
        run = h2.add_run(f"Classe {anno}ª — {AREA_LABELS.get(area, area)}")
        run.font.color.rgb = RGBColor.from_string(AREA_COLORS.get(area, "1A1A1A"))

        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Light List Accent 1"
        hdr = tbl.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Titolo MC"
        hdr[2].text = "Imm."
        for c in hdr:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(9)
            set_cell_shading(c, "EEEEEE")

        for b in sorted(by_anno_area[(anno, area)], key=lambda x: x["mc_id"]):
            row = tbl.add_row().cells
            row[0].text = b["mc_id"]
            row[1].text = b.get("mc_titolo", "")
            row[2].text = str(len(b.get("images", [])))
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ────────────────────────────────────────────────────────────────────
# MAIN
# ────────────────────────────────────────────────────────────────────

def main():
    here = Path(__file__).resolve().parent
    out_path = here / OUT_FILE

    # Trova tutti i brief .md
    files = sorted(here.glob("MC-*_visual_brief.md"))
    if not files:
        print("Nessun brief .md trovato in", here, file=sys.stderr)
        sys.exit(1)

    briefs = []
    for f in files:
        try:
            briefs.append(parse_brief_md(f.read_text(encoding="utf-8")))
        except Exception as e:
            print(f"[WARN] {f.name}: parse error {e}")

    n_brief = len(briefs)
    n_img = sum(len(b.get("images", [])) for b in briefs)
    print(f"Brief letti: {n_brief}, immagini totali: {n_img}")

    doc = Document()
    setup_styles(doc)

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    build_cover(doc, n_brief, n_img)
    build_index(doc, briefs)

    # Sezioni per ogni MC con page break tra una e l'altra
    for i, b in enumerate(briefs):
        area = re.match(r"MC-([A-Z]+)-", b.get("mc_id", "")).group(1) if b.get("mc_id") else "MAT"
        color = AREA_COLORS.get(area, "1A1A1A")
        if i > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        add_brief_section(doc, b, color)

    doc.save(out_path)
    size_kb = out_path.stat().st_size / 1024
    print(f"Scritto {out_path.name} ({size_kb:.0f} KB)")


if __name__ == "__main__":
    main()
